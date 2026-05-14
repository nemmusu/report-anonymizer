"""Round-trip integration tests for the DOCX image-redaction path."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest

docx = pytest.importorskip("docx")
PIL = pytest.importorskip("PIL")
from docx import Document
from docx.shared import Inches
from PIL import Image

from anonymize.format_adapters.docx_adapter import DocxAdapter
from anonymize.image_inventory import (
    ImageDecision,
    RedactionRect,
    compute_image_id,
)


def _png_solid(color: tuple[int, int, int], size: tuple[int, int] = (200, 200)) -> bytes:
    img = Image.new("RGB", size, color)
    buf = BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_docx(path: Path, image_bytes: bytes, *, copies: int = 1) -> None:
    doc = Document()
    doc.add_paragraph("First paragraph")
    for _ in range(copies):
        doc.add_picture(BytesIO(image_bytes), width=Inches(2.0))
    doc.save(str(path))


def test_inventory_lists_one_image(tmp_path: Path) -> None:
    src = _png_solid((220, 30, 30))
    p = tmp_path / "doc.docx"
    _make_docx(p, src)

    inv = DocxAdapter().inventory_images(p)
    assert len(inv) == 1
    item = inv[0]
    assert item.location["kind"] == "docx"
    assert item.location["rel_id"]
    assert item.fmt == "png"
    assert item.width == 200 and item.height == 200


def test_skip_decision_keeps_blob_identical(tmp_path: Path) -> None:
    src = _png_solid((10, 200, 10))
    p = tmp_path / "skip.docx"
    _make_docx(p, src)

    adapter = DocxAdapter()
    inv = adapter.inventory_images(p)
    image_id = compute_image_id(inv[0].raw_bytes)

    report = adapter.apply_image_redactions(
        p, {image_id: ImageDecision(image_id=image_id, decision="skip")}
    )
    assert report.applied == 0
    assert report.skipped == 1
    inv2 = DocxAdapter().inventory_images(p)
    # Bytes preserved exactly.
    assert compute_image_id(inv2[0].raw_bytes) == image_id


def test_redact_blackout_changes_blob_inside_rect_only(tmp_path: Path) -> None:
    src = _png_solid((220, 30, 30))
    p = tmp_path / "redact.docx"
    _make_docx(p, src)

    adapter = DocxAdapter()
    inv = adapter.inventory_images(p)
    image_id = compute_image_id(inv[0].raw_bytes)

    report = adapter.apply_image_redactions(p, {
        image_id: ImageDecision(
            image_id=image_id,
            decision="redact",
            rects=[RedactionRect(x=50, y=50, w=100, h=100, tool="blackout")],
        )
    })
    assert report.applied == 1

    inv2 = DocxAdapter().inventory_images(p)
    img = Image.open(BytesIO(inv2[0].raw_bytes))
    img.load()
    assert img.size == (200, 200)
    assert img.getpixel((100, 100)) == (0, 0, 0)
    assert img.getpixel((10, 10)) == (220, 30, 30)


def test_no_decisions_is_no_op(tmp_path: Path) -> None:
    src = _png_solid((50, 50, 50))
    p = tmp_path / "noop.docx"
    _make_docx(p, src)
    before = p.read_bytes()
    DocxAdapter().apply_image_redactions(p, {})
    after = p.read_bytes()
    # docx is a zip; rebuilds may differ in zip metadata even when
    # content is identical, so we compare via inventory rather than
    # raw bytes (the inventory walks the actual image parts).
    assert before == after


def test_logo_repeated_three_times_redacted_once(tmp_path: Path) -> None:
    src = _png_solid((30, 200, 30))
    p = tmp_path / "logo.docx"
    _make_docx(p, src, copies=3)

    adapter = DocxAdapter()
    inv = adapter.inventory_images(p)
    # python-docx may emit one ImagePart and reuse it across all
    # three add_picture calls (rel_id reused), or distinct parts.
    # Either way, all of them carry IDENTICAL bytes, so collapse by
    # image_id.
    ids = {compute_image_id(im.raw_bytes) for im in inv}
    assert len(ids) == 1
    image_id = next(iter(ids))

    adapter.apply_image_redactions(p, {
        image_id: ImageDecision(
            image_id=image_id,
            decision="redact",
            rects=[RedactionRect(x=0, y=0, w=200, h=50, tool="blackout")],
        )
    })
    # Every occurrence rewritten because there's a single ImagePart
    # behind all three; inspect inventory after.
    inv2 = DocxAdapter().inventory_images(p)
    for item in inv2:
        img = Image.open(BytesIO(item.raw_bytes))
        img.load()
        # Top band must be black for every occurrence.
        assert img.getpixel((100, 5)) == (0, 0, 0)
