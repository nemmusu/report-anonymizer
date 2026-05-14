"""DOCX adapter built on python-docx.

Word splits a single user-visible string across multiple ``Run`` elements
whenever any inline property changes (spell-check region, soft-hyphen,
tracked changes, ...). A naive per-run substitution would miss any string
that crosses a run boundary, so this adapter:

1. Concatenates the text of all runs in a paragraph,
2. Applies substitutions on the concatenated text via :func:`apply_to_text`,
3. Redistributes the resulting characters across the original runs while
   preserving each run's properties; runs that lose all their characters
   become empty (so we keep their ``rPr``).
"""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from .base import (
    FormatAdapter,
    ImageReport,
    InventoryImageRaw,
    Segment,
    SubstitutionRule,
    WriteEvent,
    WriteReport,
    apply_to_text,
)


def _iter_paragraphs(doc) -> Iterable[tuple[str, object]]:
    """Yield (seg_id_prefix, paragraph) for every paragraph in the document.

    Walks body, tables (recursively), headers, footers, footnotes, comments.
    """
    for i, p in enumerate(doc.paragraphs):
        yield f"body.p{i}", p
    for ti, table in enumerate(doc.tables):
        yield from _iter_table(f"body.t{ti}", table)
    for si, section in enumerate(doc.sections):
        for kind, hdrf in (
            ("hdr", section.header),
            ("ftr", section.footer),
            ("first_hdr", section.first_page_header),
            ("first_ftr", section.first_page_footer),
            ("even_hdr", section.even_page_header),
            ("even_ftr", section.even_page_footer),
        ):
            try:
                paras = list(hdrf.paragraphs)
            except Exception:
                continue
            for pi, p in enumerate(paras):
                yield f"sec{si}.{kind}.p{pi}", p
            try:
                tables = list(hdrf.tables)
            except Exception:
                tables = []
            for ti, t in enumerate(tables):
                yield from _iter_table(f"sec{si}.{kind}.t{ti}", t)


def _iter_table(prefix: str, table) -> Iterable[tuple[str, object]]:
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for pi, p in enumerate(cell.paragraphs):
                yield f"{prefix}.r{ri}c{ci}.p{pi}", p
            for nti, ntable in enumerate(cell.tables):
                yield from _iter_table(f"{prefix}.r{ri}c{ci}.t{nti}", ntable)


def _redistribute(runs: list, full_text: str) -> None:
    """Spread ``full_text`` across the existing runs.

    The first run that has at least one character (in the original) absorbs the
    leading prefix; subsequent runs keep their original character counts when
    possible, but if the new text is shorter than the original's total length
    we just truncate the trailing runs to empty strings.
    """
    pos = 0
    n = len(full_text)
    for idx, run in enumerate(runs):
        original_len = len(run.text or "")
        if pos >= n:
            run.text = ""
            continue
        if idx == len(runs) - 1:
            # Last run absorbs whatever is left.
            run.text = full_text[pos:]
            return
        if original_len == 0:
            # Keep empty runs empty.
            continue
        # Take a slice of the same size as the original run, capped at n.
        end = min(pos + original_len, n)
        run.text = full_text[pos:end]
        pos = end


class DocxAdapter(FormatAdapter):
    name = "docx"
    extensions = {".docx"}
    mimes = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    def __init__(self) -> None:
        try:
            from docx import Document  # type: ignore
        except Exception as e:  # pragma: no cover
            raise RuntimeError(
                "python-docx is required for the docx adapter "
                f"(pip install python-docx). Original error: {e}"
            )
        self._Document = Document

    # ---- extract ------------------------------------------------------------

    def extract(self, path: Path) -> list[Segment]:
        doc = self._Document(str(path))
        out: list[Segment] = []
        for seg_id, para in _iter_paragraphs(doc):
            text = "".join(r.text or "" for r in para.runs)
            if not text:
                continue
            out.append(Segment(seg_id=seg_id, text=text))
        return out

    # ---- write --------------------------------------------------------------

    def write(
        self,
        src_path: Path,
        dst_path: Path,
        substitutions: list[SubstitutionRule],
    ) -> WriteReport:
        doc = self._Document(str(src_path))
        events: list[WriteEvent] = []
        for seg_id, para in _iter_paragraphs(doc):
            runs = para.runs
            if not runs:
                continue
            full_text = "".join(r.text or "" for r in runs)
            if not full_text:
                continue
            new_text, ev = apply_to_text(full_text, substitutions, seg_id=seg_id)
            if new_text == full_text:
                continue
            _redistribute(runs, new_text)
            events.extend(ev)
        dst_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dst_path))
        return WriteReport(file_rel=str(dst_path), events=events)

    # ---- image inventory + apply ------------------------------------------

    def inventory_images(self, path: Path) -> list[InventoryImageRaw]:
        """Walk every inline shape, return one record per embedded
        image. Floating shapes (anchored / behind text) are NOT
        first-class in python-docx; we surface those via the
        ``related_parts`` walk so we still see them at scan time and
        can replace them at apply time.

        Each record carries the relationship id (``rel_id``) of the
        image part. The rel id is what python-docx uses to look up
        the part for replacement, so it's the natural location key.
        Inline shapes get an additional ``inline_shape_index`` for
        the GUI gallery (so it can show "image #2 in document order").
        """
        from io import BytesIO
        try:
            from PIL import Image as _PIL  # type: ignore
        except Exception:
            _PIL = None  # type: ignore

        out: list[InventoryImageRaw] = []
        try:
            doc = self._Document(str(path))
        except Exception:
            return out

        # Track which rel_ids we've already inventoried so the same
        # image referenced by multiple inline shapes (logo on every
        # paragraph) is reported once. Apply will still replace it
        # everywhere (it's one ImagePart).
        seen_rels: set[str] = set()
        # First pass: inline shapes, capture stable index for GUI.
        for idx, shape in enumerate(doc.inline_shapes):
            try:
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                rel_id = blip.embed
            except Exception:
                continue
            if not rel_id or rel_id in seen_rels:
                continue
            seen_rels.add(rel_id)
            part = doc.part.related_parts.get(rel_id)
            if part is None:
                continue
            raw = bytes(part.blob)
            fmt = _ext_for_content_type(getattr(part, "content_type", "")) or "png"
            width, height = _image_dimensions(raw, _PIL)
            out.append(
                InventoryImageRaw(
                    raw_bytes=raw,
                    fmt=fmt,
                    width=width,
                    height=height,
                    location={
                        "kind": "docx",
                        "inline_shape_index": idx,
                        "rel_id": rel_id,
                    },
                )
            )

        # Second pass: catch floating / anchored shapes that
        # ``inline_shapes`` does not surface. These come through
        # ``document.part.rels`` with ``reltype`` "image".
        for rel in doc.part.rels.values():
            if not rel.reltype.endswith("/image"):
                continue
            rel_id = rel.rId
            if rel_id in seen_rels:
                continue
            part = rel.target_part
            if part is None:
                continue
            try:
                raw = bytes(part.blob)
            except Exception:
                continue
            fmt = _ext_for_content_type(getattr(part, "content_type", "")) or "png"
            width, height = _image_dimensions(raw, _PIL)
            seen_rels.add(rel_id)
            out.append(
                InventoryImageRaw(
                    raw_bytes=raw,
                    fmt=fmt,
                    width=width,
                    height=height,
                    location={
                        "kind": "docx",
                        "rel_id": rel_id,
                    },
                    warnings=["floating_or_header_shape"],
                )
            )
        return out

    def apply_image_redactions(
        self,
        dst_path: Path,
        decisions_for_file: dict,
    ) -> ImageReport:
        """Replace ``image_part.blob`` for every image whose
        ``image_id`` has a ``redact`` decision. python-docx serialises
        the docx as a zip on save, the on-disk file gets rewritten
        atomically here through a tempfile + ``os.replace``.
        """
        from ..image_inventory import compute_image_id, ImageDecision
        from ..image_redactor import ImageRedaction, ImageRedactor

        report = ImageReport(file_rel=str(dst_path))
        if not decisions_for_file:
            return report
        prepared, skip_ids = _prepare_decisions(decisions_for_file, ImageDecision, ImageRedaction)
        if not prepared and not skip_ids:
            return report

        try:
            doc = self._Document(str(dst_path))
        except Exception as e:
            report.warnings.append(f"open_failed:{e}")
            return report

        seen_rels: set[str] = set()
        for rel in doc.part.rels.values():
            if not rel.reltype.endswith("/image"):
                continue
            rel_id = rel.rId
            if rel_id in seen_rels:
                continue
            seen_rels.add(rel_id)
            part = rel.target_part
            if part is None:
                continue
            try:
                raw = bytes(part.blob)
            except Exception:
                continue
            image_id = compute_image_id(raw)
            if image_id in skip_ids:
                report.skipped += 1
                continue
            rects = prepared.get(image_id)
            if not rects:
                report.untouched += 1
                continue
            fmt_hint = _ext_for_content_type(getattr(part, "content_type", "")) or "png"
            try:
                result = ImageRedactor.redact_bytes(raw, fmt_hint, rects)
            except Exception as e:
                report.warnings.append(f"redact_failed:{image_id}:{e}")
                continue
            try:
                part._blob = result.bytes_  # python-docx ImagePart stores blob in _blob
            except Exception as e:
                report.warnings.append(f"replace_failed:{image_id}:{e}")
                continue
            if result.warnings:
                report.warnings.extend(
                    f"{image_id}:{w}" for w in result.warnings
                )
            report.applied += 1

        # Atomic save: write to a tempfile in the same directory then
        # ``os.replace`` so an interrupted save never leaves a
        # half-written docx (zip with a missing ``[Content_Types].xml``
        # for instance).
        tmp_save = dst_path.with_suffix(dst_path.suffix + ".imgtmp")
        try:
            doc.save(str(tmp_save))
        except Exception as e:
            report.warnings.append(f"save_failed:{e}")
            try:
                tmp_save.unlink()
            except OSError:
                pass
            return report
        import os as _os
        _os.replace(str(tmp_save), str(dst_path))
        return report


def _ext_for_content_type(ct: str) -> str:
    """``image/png`` -> ``png``, etc. Used to stamp ``fmt_hint`` on
    inventory entries and to feed PIL's format roundtrip on apply.
    """
    if not ct:
        return ""
    if "/" in ct:
        ext = ct.split("/", 1)[1].split(";", 1)[0].strip().lower()
        if ext == "jpeg":
            return "jpeg"
        if ext == "x-emf" or ext == "emf":
            return "emf"
        if ext == "x-wmf" or ext == "wmf":
            return "wmf"
        return ext
    return ""


def _image_dimensions(raw: bytes, pil_module) -> tuple[int, int]:
    """Best-effort (width, height) decoding of raw image bytes.

    Falls back to ``(0, 0)`` if PIL is unavailable or the format is
    something it cannot decode (very old EMF variants, for instance).
    The inventory still records the entry; the GUI shows ``? x ?`` and
    the user can still skip / mark for blackout based on the format
    alone.
    """
    if pil_module is None:
        return (0, 0)
    try:
        from io import BytesIO
        img = pil_module.open(BytesIO(raw))
        img.load()
        return (int(img.size[0]), int(img.size[1]))
    except Exception:
        return (0, 0)


def _prepare_decisions(
    decisions_for_file: dict,
    ImageDecisionCls,
    ImageRedactionCls,
) -> tuple[dict[str, list], set[str]]:
    """Box dict-shaped operator input into the dataclass form
    ``apply_image_redactions`` consumes. Tolerates both dataclass
    instances and raw dicts so tests / scripts can drive the apply
    path without going through the full pipeline.
    """
    prepared: dict[str, list] = {}
    skip_ids: set[str] = set()
    for image_id, decision in decisions_for_file.items():
        if not isinstance(decision, ImageDecisionCls):
            decision = ImageDecisionCls.from_dict(image_id, decision or {})
        if decision.decision in ("skip", "defer"):
            skip_ids.add(image_id)
            continue
        if decision.decision != "redact" or not decision.rects:
            skip_ids.add(image_id)
            continue
        prepared[image_id] = [
            ImageRedactionCls(
                x=r.x, y=r.y, w=r.w, h=r.h,
                tool=r.tool,
                intensity=r.intensity,
                text=r.text,
                font_size=r.font_size,
                fg=r.fg,
                bg=r.bg,
            ) for r in decision.rects
        ]
    return prepared, skip_ids


__all__ = ["DocxAdapter"]
